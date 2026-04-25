# import fitz  # PyMuPDF
# from docx import Document


# def extract_text_from_pdf(file_path):
#     text = ""
#     doc = fitz.open(file_path)

#     for page in doc:
#         text += page.get_text()

#     return text


# def extract_text_from_docx(file_path):
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])


# def extract_text_from_txt(file_path):
#     with open(file_path, "r", encoding="utf-8") as f:
#         return f.read()
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO


def extract_text_from_pdf(file_bytes):
    text = ""

    doc = fitz.open(stream=file_bytes, filetype="pdf")

    for page in doc:
        text += page.get_text()

    return text


def extract_text_from_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_txt(file_bytes):
    return file_bytes.decode("utf-8")